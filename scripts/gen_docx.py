#!/data/data/com.termux/files/home/.hermes/venv/bin/python3
"""Generate Verano Media Brand Guidelines as .docx"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Page Setup ───
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Style helpers ───
BLUE = RGBColor(0, 102, 204)
DARK = RGBColor(28, 28, 30)
GRAY = RGBColor(142, 142, 147)
WHITE = RGBColor(255, 255, 255)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = BLUE
            run.font.size = Pt(26)
        elif level == 2:
            run.font.color.rgb = DARK
            run.font.size = Pt(18)
        elif level == 3:
            run.font.color.rgb = DARK
            run.font.size = Pt(14)
    return h

def add_para(text, bold=False, color=DARK, size=11, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, bold_prefix="", color=DARK):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = color
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.color.rgb = color
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.color.rgb = color
        run.font.size = Pt(11)
    return p

def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '0066CC')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.bold = True
                run.font.size = Pt(10)
    
    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F5F7')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = DARK
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    
    return table

# ═══════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Verano Media")
run.font.size = Pt(48)
run.font.color.rgb = BLUE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Manual de Identidad Visual")
run.font.size = Pt(24)
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Tu marca en su mejor temporada"')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Agencia de Marketing Digital")
run.font.size = Pt(14)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Versión 1.0 — Mayo 2026")
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_page_break()

# ═══════════════════════════════════════
# TABLE OF CONTENTS (simple)
# ═══════════════════════════════════════
add_heading("Contenido", level=1)
toc_items = [
    "1. Marca",
    "2. Logotipo (Wordmark)",
    "3. Isotipo / Símbolo",
    "4. Paleta de Color",
    "5. Tipografía",
    "6. Elementos Gráficos",
    "7. Fotografía e Imágenes",
    "8. Tono de Voz",
    "9. Aplicaciones en Redes y Web",
    "10. Checklist de Consistencia"
]
for item in toc_items:
    add_para(item, size=12)

doc.add_page_break()

# ═══════════════════════════════════════
# 1. MARCA
# ═══════════════════════════════════════
add_heading("1. Marca", level=1)

add_heading("Nombre Completo", level=2)
add_para("Verano Media — todo junto, con V mayúscula y M mayúscula (camelCase).")

add_heading("Tagline / Claim", level=2)
add_para('"Tu marca en su mejor temporada"', italic=True, size=13, color=BLUE)

add_heading("Descripción Corporativa", level=2)
add_para("Agencia de Marketing Digital. No usar traducciones al inglés ni variantes como 'Agencia de Marketing Digital en RD' a menos que sea necesario por contexto.")

add_heading("Esencia de Marca", level=2)
add_bullet("Somos dominicanos, profesionales y cercanos")
add_bullet("Combinamos estrategia, diseño y tecnología")
add_bullet("Enseñamos mientras vendemos")
add_bullet("No prometemos lo que no podemos cumplir")

doc.add_page_break()

# ═══════════════════════════════════════
# 2. LOGOTIPO
# ═══════════════════════════════════════
add_heading("2. Logotipo (Wordmark)", level=1)

add_heading("Estructura", level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("VeranoMedia")
run.font.size = Pt(28)
run.font.color.rgb = DARK
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AGENCIA DE MARKETING DIGITAL")
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.name = 'Inter'

add_heading("Colores del Logotipo", level=2)
add_table(
    ["Elemento", "Color", "Hex", "RGB", "Uso"],
    [
        ["Verano", "Gris carbón", "#1C1C1E", "(28, 28, 30)", "Nombre de marca — parte sólida"],
        ["Media", "Azul marca", "#0066CC", "(0, 102, 204)", "Nombre de marca — parte digital"],
        ["Tagline", "Gris medio", "#8E8E93", "(142, 142, 147)", "Descriptor secundario"],
        ["Línea decorativa", "Azul marca", "#0066CC", "(0, 102, 204)", "Debajo de 'Media'"],
    ],
    col_widths=[1.2, 1.2, 1, 1.2, 2.5]
)

add_heading("Tipografía del Logotipo", level=2)
add_para("Sans-serif humanista / neo-grotesca. Compatible con Inter, DM Sans o Plus Jakarta Sans.")
add_bullet("'Verano': peso semibold a bold, color #1C1C1E")
add_bullet("'Media': peso semibold a bold, color #0066CC")
add_bullet("'Media' lleva línea horizontal fina debajo del mismo azul")
add_bullet("Tagline: peso light/regular, mayúsculas sostenidas, tracking amplio")

add_heading("Variantes del Logotipo", level=2)
add_table(
    ["Variante", "Descripción", "Cuándo usarla"],
    [
        ["Completo", "Verano + Media + Tagline", "Uso estándar, web, documentos"],
        ["Simplificado", "Solo 'VeranoMedia' sin tagline", "Espacios reducidos, redes"],
        ["Monograma", "'VM' en blanco sobre fondo azul", "Favicon, avatar, app icon"],
        ["Versión oscura", "Logotipo en blanco", "Sobre fondos oscuros/azules"],
        ["Versión clara", "Logotipo en gris oscuro", "Sobre fondos claros"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

add_heading("Área de Resguardo", level=2)
add_para("Mantener un margen mínimo del 10% del ancho del logotipo alrededor del mismo. No colocar texto, iconos u otros elementos dentro de esta área.")

add_heading("Tamaño Mínimo", level=2)
add_bullet("Digital: 120px (completo) / 60px (solo wordmark) / 32px (monograma)")
add_bullet("Impresión: 25mm (completo) / 12mm (monograma)")

add_heading("Usos Incorrectos ❌", level=2)
add_bullet("No separar 'Verano' y 'Media' en líneas independientes sin alineación")
add_bullet("No cambiar el color de 'Verano' (siempre #1C1C1E)")
add_bullet("No cambiar el color de 'Media' (siempre #0066CC)")
add_bullet("No rotar, distorsionar, estirar ni comprimir el logotipo")
add_bullet("No aplicar sombras excesivas o efectos 3D no aprobados")
add_bullet("No colocar sobre fondos que comprometan la legibilidad")
add_bullet("No usar 'VERANOMEDIA' todo en mayúsculas")

doc.add_page_break()

# ═══════════════════════════════════════
# 3. ISOTIPO
# ═══════════════════════════════════════
add_heading("3. Isotipo / Símbolo", level=1)

add_heading("Monograma VM", level=2)
add_para("'VM' en blanco, tipografía sans-serif bold, sobre fondo azul oscuro (#0A1628) o degradado azul marino. Usado como favicon, app icon y avatar pequeño.")

add_heading("Opciones de Isotipo", level=2)
add_bullet("A — VM entrelazado: iniciales combinadas en un símbolo limpio y moderno")
add_bullet("B — Sol geométrico: círculo con rayos estilizados, conecta con 'Verano'")
add_bullet("C — Sparkle/estrella de 4 puntas: el elemento decorativo del brand board")
add_bullet("D — VM dentro de círculo: estilo app icon, blanco sobre fondo azul")

add_heading("Detalle Decorativo", level=2)
add_para("Estrella de 4 puntas (sparkle) que aparece en el brand board. NO es parte obligatoria del logotipo, pero puede usarse como detalle decorativo en piezas gráficas, separadores visuales o elementos de cierre.")

doc.add_page_break()

# ═══════════════════════════════════════
# 4. PALETA DE COLOR
# ═══════════════════════════════════════
add_heading("4. Paleta de Color", level=1)

add_heading("Colores Principales", level=2)
add_table(
    ["Nombre", "Hex", "RGB", "Uso Principal"],
    [
        ["Azul Marca (logo)", "#0066CC", "(0, 102, 204)", "Logotipo 'Media', acentos visuales"],
        ["Azul Primario Digital", "#007AFF", "(0, 122, 255)", "Botones CTAs, links, hover, selección"],
        ["Gris Carbón", "#1C1C1E", "(28, 28, 30)", "Texto principal, color 'Verano' del logo"],
        ["Blanco", "#FFFFFF", "(255, 255, 255)", "Fondos principales, texto sobre oscuro"],
    ],
    col_widths=[1.8, 1.2, 1.2, 2.8]
)

add_para("NOTA: El azul del logo (#0066CC) y el azul digital (#007AFF) son distintos. #0066CC = logo y acentos. #007AFF = elementos interactivos. Esta dualidad es intencional.", bold=True, size=10, color=BLUE)

add_heading("Neutros (Escala de Grises)", level=2)
add_table(
    ["Nombre", "Hex", "RGB", "Uso"],
    [
        ["Fondo claro", "#F5F5F7", "(245, 245, 247)", "Secciones alternas, fondos de cards"],
        ["Borde sutil", "#E8E8ED", "(232, 232, 237)", "Bordes internos, separadores"],
        ["Borde medio", "#D1D1D6", "(209, 209, 214)", "Inputs, bordes de tarjetas"],
        ["Texto secundario", "#8E8E93", "(142, 142, 147)", "Tagline, metadata, descripciones"],
        ["Texto terciario", "#636366", "(99, 99, 102)", "Notas al pie, texto pequeño"],
    ],
    col_widths=[1.5, 1.2, 1.2, 3]
)

add_heading("Colores Semánticos", level=2)
add_table(
    ["Nombre", "Hex", "Uso"],
    [
        ["Success ✅", "#34C759", "Éxito, completado, checkmarks"],
        ["Warning ⚠️", "#FF9F0A", "Advertencias, atención"],
        ["Error ❌", "#FF453A", "Errores, cancelaciones"],
        ["Info ℹ️", "#5AC8FA", "Información, notificaciones"],
    ],
    col_widths=[2, 1.5, 3.5]
)

add_heading("Proporciones de Color en el Sitio", level=2)
add_bullet("Blanco (#FFFFFF): ~65% (fondos principales)")
add_bullet("Gris carbón (#1C1C1E): ~15% (texto)")
add_bullet("Fondo claro (#F5F5F7): ~10% (fondos alternos)")
add_bullet("Azules (#0066CC / #007AFF): ~7% (acentos, botones)")
add_bullet("Otros (neutros + semánticos): ~3%")

doc.add_page_break()

# ═══════════════════════════════════════
# 5. TIPOGRAFÍA
# ═══════════════════════════════════════
add_heading("5. Tipografía", level=1)

add_heading("Fuentes Principales", level=2)
add_table(
    ["Función", "Fuente", "Pesos", "Uso"],
    [
        ["Sans-serif (primaria)", "Inter", "400, 500, 600, 700, 800", "Navegación, body, cards, precios"],
        ["Serif (editorial)", "Playfair Display", "600, 700", "Títulos de blog, contenido editorial"],
    ],
    col_widths=[2, 2, 2, 2]
)

add_heading("Jerarquía Tipográfica", level=2)
add_table(
    ["Token", "Tamaño", "Peso", "Letter-spacing", "Uso"],
    [
        ["display", "5rem / 80px", "Extrabold (800)", "-0.04em", "Hero principal"],
        ["heading-xl", "4rem / 64px", "Bold (700)", "-0.03em", "Títulos de página"],
        ["heading-lg", "3rem / 48px", "Bold (700)", "-0.02em", "Títulos de sección"],
        ["heading-md", "2rem / 32px", "Semibold (600)", "-0.01em", "Títulos de tarjetas"],
        ["heading-sm", "1.5rem / 24px", "Semibold (600)", "-0.01em", "Subtítulos"],
        ["body", "1rem / 16px", "Regular (400)", "normal", "Párrafos, contenido general"],
        ["small", "0.875rem / 14px", "Regular/Medium", "normal", "Etiquetas, metadatos"],
        ["caption", "0.75rem / 12px", "Medium (500)", "normal", "Notas al pie, créditos"],
    ],
    col_widths=[1.3, 1.5, 1.5, 1.3, 2.5]
)

add_heading("Reglas de Blog (Estilo Editorial)", level=2)
add_bullet("Títulos: Playfair Display Bold, color #1C1C1E")
add_bullet("Body: Inter Regular 18px, color #1C1C1E, line-height 1.7")
add_bullet("Drop cap: primera letra en Playfair Display ~64px, color #0066CC")
add_bullet("Pull quotes: Playfair Display italic, barra lateral azul")
add_bullet("Data boxes: fondo #F5F5F7, border-left 4px #0066CC")
add_bullet("Sin imágenes en hero de blog — solo tipografía pura")

doc.add_page_break()

# ═══════════════════════════════════════
# 6. ELEMENTOS GRÁFICOS
# ═══════════════════════════════════════
add_heading("6. Elementos Gráficos", level=1)

add_heading("Bordes Redondeados", level=2)
add_table(
    ["Token", "Valor", "Uso"],
    [
        ["vm-sm", "6px", "Botones pequeños, inputs"],
        ["vm-md", "12px", "Botones principales, cards"],
        ["vm-lg", "20px", "Tarjetas destacadas"],
        ["vm-xl", "32px", "Modales, contenedores hero"],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

add_heading("Sombras", level=2)
add_table(
    ["Token", "Valor"],
    [
        ["vm-sm", "0px 1px 2px rgba(28, 28, 30, 0.05)"],
        ["vm-md", "0px 4px 12px rgba(28, 28, 30, 0.08)"],
        ["vm-lg", "0px 8px 30px rgba(28, 28, 30, 0.12)"],
        ["vm-xl", "0px 20px 60px rgba(28, 28, 30, 0.15)"],
    ],
    col_widths=[1.5, 5]
)

add_heading("Layout", level=2)
add_bullet("Contenedor máximo: 1200px")
add_bullet("Padding lateral: 1.5rem (mobile) / 2rem (desktop)")
add_bullet("Separación entre secciones: 4rem")
add_bullet("Fondo alterno de sección: #F5F5F7")

doc.add_page_break()

# ═══════════════════════════════════════
# 7. FOTOGRAFÍA
# ═══════════════════════════════════════
add_heading("7. Fotografía e Imágenes", level=1)

add_heading("Estilo Visual", level=2)
add_bullet("Fotos de equipo: reales, profesionales, con personalidad — nada de stock photos")
add_bullet("Hero del sitio: video loop 848×478, 10s, mano robótica escribiendo + marca")
add_bullet("Poster del video: hero-frame.jpg (fondo oscuro/premium)")
add_bullet("Sin filtros pesados — colores naturales, bien iluminados")

add_heading("Fotos de Miembros del Equipo", level=2)
add_bullet("Formato: JPG, rostro visible, fondo neutro o contextual")
add_bullet("Tamaño en web: ~300×300px, recorte cuadrado")
add_bullet("Estilo: profesional pero relajado, cada uno con su personalidad")

doc.add_page_break()

# ═══════════════════════════════════════
# 8. TONO DE VOZ
# ═══════════════════════════════════════
add_heading("8. Tono de Voz", level=1)

add_heading("Personalidad de Marca", level=2)
add_table(
    ["Atributo", "Descripción"],
    [
        ["Profesional", "Sabemos lo que hacemos, lo decimos con seguridad"],
        ["Cálido", "Somos caribeños dominicanos, se nota en el trato"],
        ["Directo", "Sin rodeos, sin jerga innecesaria"],
        ["Educativo", "Enseñamos mientras vendemos (blog y LinkedIn)"],
        ["Confiable", "No prometemos lo que no podemos cumplir"],
    ],
    col_widths=[2, 4.5]
)

add_heading("Frases Clave de Marca", level=2)
add_bullet('Claim principal: "Tu marca en su mejor temporada"')
add_bullet('Propuesta de valor: "Diseñamos la presencia digital que tu negocio necesita"')
add_bullet('Lead magnet: "Diagnóstico Digital Gratuito"')
add_bullet('Footer / CTA: "Sin compromiso, sin spam. Solo valor real para tu negocio."')

add_heading("Reglas de Idioma", level=2)
add_bullet("Español latinoamericano neutro con toque dominicano")
add_bullet("No usar spanglish forzado")
add_bullet("Términos técnicos en inglés solo cuando no hay traducción clara (SEO, ROI, landing page)")

doc.add_page_break()

# ═══════════════════════════════════════
# 9. APLICACIONES
# ═══════════════════════════════════════
add_heading("9. Aplicaciones en Redes y Web", level=1)

add_heading("Redes Sociales", level=2)
add_table(
    ["Plataforma", "Foto de Perfil", "Bio"],
    [
        ["LinkedIn", "Logo completo + tagline (fondo azul marino)", "Agencia de Marketing Digital. Ayudamos a marcas y negocios a crecer con estrategia, diseño y tecnología."],
        ["Instagram", "Monograma VM o logo simplificado", "Agencia de Marketing Digital • RD • Estrategia • Diseño Web • Redes Sociales • SEO • Automatización"],
        ["Facebook", "Logo completo (fondo azul marino)", "Agencia de Marketing Digital con enfoque en resultados. Diseño web, redes sociales, SEO, automatización y más."],
    ],
    col_widths=[1.3, 2.5, 3.5]
)

add_heading("Website", level=2)
add_bullet("URL: veranomedia.digital (pendiente de configuración)")
add_bullet("Deploy actual: GitHub Pages (peterfoxx93-design.github.io/veranomedia)")
add_bullet("Stack: React + Vite + Tailwind CSS + Motion")
add_bullet("Páginas: Inicio, Servicios (10+), Portafolio, Blog, Nosotros, Contacto")

add_heading("Email Corporativo", level=2)
add_para("Pendiente de configuración. Formato sugerido: nombre@veranomedia.digital")
add_bullet("Principal: hola@veranomedia.digital")
add_bullet("Alternativo: info@veranomedia.digital")

doc.add_page_break()

# ═══════════════════════════════════════
# 10. CHECKLIST
# ═══════════════════════════════════════
add_heading("10. Checklist de Consistencia de Marca", level=1)

add_para("Antes de publicar cualquier pieza, verificar:")

checklist = [
    "¿'Verano' está en gris carbón (#1C1C1E) y 'Media' en azul (#0066CC)?",
    "¿El logotipo tiene el área de resguardo adecuada?",
    "¿Los colores usados corresponden a la paleta oficial?",
    "¿La tipografía es Inter (body) o Playfair Display (blog)?",
    "¿El tono es profesional pero cálido?",
    "¿El claim 'Tu marca en su mejor temporada' está presente donde corresponde?",
    "¿Las imágenes son reales (no stock photos genéricas)?",
    "¿El CTA de WhatsApp tiene el número correcto (1-829-684-8477)?",
    "¿Los enlaces a redes sociales dirigen a las cuentas correctas?",
]

for i, item in enumerate(checklist, 1):
    add_para(f"☐  {i}. {item}", size=11)

# ─── Footer ───
doc.add_paragraph()
add_para("—", color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Documento generado: 13 de mayo de 2026", color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Versión 1.0", color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Próxima actualización: al definir el isotipo final y las redes sociales", color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# ─── Save ───
output_path = "/data/data/com.termux/files/home/veranomedia/public/Manual-Identidad-Visual-VeranoMedia.docx"
doc.save(output_path)
print(f"✅ Documento guardado: {output_path}")
print(f"   Tamaño: {os.path.getsize(output_path) / 1024:.1f} KB")
